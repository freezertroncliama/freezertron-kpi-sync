"""
KPI de Manutenção — múltiplos clientes (Mangels S/A + Noel Supermercados)
-----------------------------------------------------
Varre recursivamente os relatórios em PDF (gerados pelo app Produttivo) no
Nextcloud, classifica cada um como preventiva ou corretiva, extrai data e
duração, e gera um relatório .docx com KPIs por categoria e por
equipamento — um cliente por vez, configurado em CLIENTES.

Premissas validadas com dados reais (ver histórico do projeto):
- Estrutura Mangels: ANO 2026/<Categoria>/<Equipamento>/<arquivo>.pdf
  (profundidade fixa, 2 níveis sob a pasta raiz).
- Estrutura Noel: profundidade INCONSISTENTE — equipamento pode estar
  direto sob a loja (LOJA 01/N 021 - Rack de Resfriado/) ou aninhado sob
  uma pasta extra (LOJA 01/NOEL LOJA 01/NLJ1_001_BALCÃO_.../). Por isso,
  pro Noel: `equipamento` é sempre a pasta imediatamente pai do PDF (não
  depende de profundidade fixa); `loja` vem do prefixo NLJ1-4 no caminho
  (confirmado com o Rafael em 06/08/2026: NLJ1=Loja1 ... NLJ4=Loja4), com
  fallback pra "LOJA N" se não achar o prefixo; `categoria` vem de
  casamento de palavra-chave no nome da pasta de equipamento (lista
  também passada pelo Rafael) — não do texto do PDF, porque o "Local do
  ativo" registrado no Produttivo (ex: "SALA DE MÁQUINAS") não bate com o
  tipo de equipamento que o time quer ver no painel (Balcão/Câmara/Ilha/...).
- Tipo: definido pelas primeiras linhas do texto do PDF, não pelo nome do
  arquivo (o nome do arquivo pode estar errado — já vimos casos reais).
  Regra (confirmada com o Rafael em 30/07/2026):
    * título começa com "PMOC"                -> preventiva
    * título começa com "Manutenção Corretiva" -> corretiva (formulário
      novo, em uso desde jul/2026, só pra corretivas)
    * título começa com "Check list" -> vencimento_filtro (checklist de
      troca de filtro de bebedouro, não é atendimento de equipamento — ver
      extrair_datas_filtro)
    * título começa com "Ordem de Serviço":
        - categoria CHILLER'S (Mangels) -> preventiva (chillers ainda não
          têm template PMOC próprio; até criarem um, o time usa "Ordem de
          Serviço" pras visitas de rotina)
        - qualquer outra categoria -> corretiva (chamado avulso)
  Confirmado em 06/08/2026 com um PDF real do Noel: mesmo formato
  Produttivo da Mangels (mesmo cabeçalho "Em: dd/mm/aaaa hh:mm", mesmas
  seções Diagnóstico/Solução/Fim do Trabalho) — a classificação por texto
  vale sem alteração pros dois clientes.
  Cuidado: essa regra pode classificar como preventiva alguma corretiva
  de chiller registrada ANTES do formulário "Manutenção Corretiva"
  existir (pré-julho/2026) — não dá pra distinguir com o dado disponível.
- Data: campo "Em: dd/mm/aaaa hh:mm" no topo do PDF (fallback: data no
  nome do arquivo, pro caso de PDF escaneado sem texto extraível).
- Duração: só é calculada para corretivas, via campos "Início do
  Trabalho" / "Fim do Trabalho". PMOC não tem duração real de serviço
  no PDF (só tempo de preenchimento do checklist no app, que não
  representa trabalho em campo).
- Vazamento/gás: pra alimentar a seção "Recorrência de vazamento" do
  painel do Noel, marca `vazamento=True` quando a palavra aparece no
  texto do PDF, e tenta extrair uma quantidade em kg (`extrair_gas_kg`).
  Confirmado com o Rafael em 06/08/2026: os relatórios de vazamento
  ainda não têm campo estruturado de quantidade — a extração retorna
  None até o time começar a preencher isso nos relatórios; não é erro.

Depois de gerar o .docx, publica os atendimentos na tabela
kpi_manutencao_atendimentos do Supabase (mesmo projeto do gerador de
orçamento) via service role key — é assim que a aba /kpi-manutencao do
Next.js exibe os dados, sem o Next.js falar com o Nextcloud diretamente e
sem tocar em nada de Omie/cadastro de clientes. Cada cliente publica só
as próprias linhas (DELETE escopado por `cliente`) — processar um cliente
nunca apaga o outro.

USO
1. Preencher .env com NEXTCLOUD_URL, NEXTCLOUD_USER, NEXTCLOUD_APP_PASSWORD,
   SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY.
2. python kpi_manutencao.py                  # roda todos os clientes
   python kpi_manutencao.py --cliente Noel    # roda só um, pra testar isolado
"""

import json
import os
import re
import shutil
import sys
import unicodedata
from dataclasses import asdict, dataclass, field
from datetime import datetime, timedelta

import pdfplumber
import requests
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from webdav3.client import Client

# Terminal/arquivo de saída no Windows costuma usar cp1252, que não sabe
# codificar acentos nem símbolos (⚠, ✓ etc.) — força UTF-8 pra print() não
# quebrar quando a saída for redirecionada pra arquivo.
sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

load_dotenv()

NEXTCLOUD_URL = os.environ["NEXTCLOUD_URL"]
NEXTCLOUD_USER = os.environ["NEXTCLOUD_USER"]
NEXTCLOUD_APP_PASSWORD = os.environ["NEXTCLOUD_APP_PASSWORD"]
SUPABASE_URL = os.environ["SUPABASE_URL"]
SUPABASE_SERVICE_ROLE_KEY = os.environ["SUPABASE_SERVICE_ROLE_KEY"]

PASTA_TEMP = "temp_pdfs"
# Bump sempre que classificar_tipo/extrair_data/extrair_duracao_corretiva
# ou a forma do Atendimento mudar de comportamento — invalida entradas de
# cache presas na regra antiga.
CACHE_VERSAO = 8
# Periodicidade da troca de filtro de bebedouro (fixa, sem template de PMOC
# próprio) — usada só pra corrigir o bug abaixo, o cálculo real de
# conforme/vencendo/atrasado é feito no painel (Next.js).
PERIODICIDADE_FILTRO_DIAS = 90

# Um cliente por entrada: nome (vai pra coluna `cliente` no Supabase),
# pasta raiz no Nextcloud, e o cache de etags (separado por cliente pra
# não misturar o "já processei isso" de um com o do outro).
CLIENTES = [
    {
        "nome": "Mangels",
        "pasta_raiz": "REFRIGERAÇÃO AÇOS - RELATÓRIOS/ANO 2026",
        "cache_path": "cache_pdfs_mangels.json",
    },
    {
        "nome": "Noel",
        "pasta_raiz": "REFRIGERAÇÃO NOEL SUPERMERCADOS",
        "cache_path": "cache_pdfs_noel.json",
    },
]

# Palavra-chave (sem acento, maiúscula) -> categoria exibida no painel.
# Lista passada pelo Rafael em 06/08/2026, pro painel do Noel. Ordem
# importa pouco aqui (nenhuma chave é substring de outra de forma
# ambígua), mas mantém "PREPARO" depois de "AREA DE PREPARO" só por
# clareza de leitura.
CATEGORIAS_NOEL: list[tuple[str, str]] = [
    ("BALCAO", "Balcão"),
    ("CAMARA", "Câmara"),
    ("ILHA", "Ilha"),
    ("EXPOSITOR", "Expositor"),
    ("FREEZER", "Freezer"),
    ("AREA DE PREPARO", "Área de Preparo"),
    ("PREPARO", "Área de Preparo"),
    ("RACK", "Rack"),
    ("SALA DE MAQUINA", "Sala de Máquina"),
    ("WALK IN COOLER", "Walk-in Cooler"),
    ("WALKINCOOLER", "Walk-in Cooler"),
    ("CERVEJEIRA", "Cervejeira"),
    ("PURIFICADOR", "Purificador de Água"),
    ("CLIMATICA", "Climática"),
]


@dataclass
class Atendimento:
    cliente: str
    categoria: str
    equipamento: str
    arquivo: str
    tipo: str  # 'preventiva' | 'corretiva' | 'desconhecida' | 'vencimento_filtro'
    data: datetime | None
    duracao_minutos: float | None
    loja: str | None = None
    ultima_substituicao_filtro: datetime | None = None
    proxima_substituicao_filtro: datetime | None = None
    periodicidade: str | None = None  # 'semanal' | 'mensal' | 'trimestral' | None
    vazamento: bool = False
    gas_kg: float | None = None


def conectar_nextcloud() -> Client:
    options = {
        "webdav_hostname": f"{NEXTCLOUD_URL}/remote.php/dav/files/{NEXTCLOUD_USER}",
        "webdav_login": NEXTCLOUD_USER,
        "webdav_password": NEXTCLOUD_APP_PASSWORD,
    }
    return Client(options)


def listar_arvore(client: Client, pasta_raiz: str) -> list[dict]:
    """Uma chamada só (Depth: infinity) pra árvore inteira, em vez de uma
    chamada por pasta — o jeito anterior (uma requisição por pasta de
    equipamento, ~100+ pastas) é o motivo da varredura levar ~10min só pra
    listar; isso aqui deve levar segundos. Retorna cada item com o
    `path` já relativo à raiz do WebDAV (sem o prefixo /remote.php/...)."""
    prefixo_servidor = f"/remote.php/dav/files/{NEXTCLOUD_USER}/"

    def caminho_relativo(item: dict) -> str:
        caminho_servidor = item["path"]
        if caminho_servidor.startswith(prefixo_servidor):
            caminho = caminho_servidor[len(prefixo_servidor):]
        else:
            caminho = caminho_servidor.lstrip("/")
        return caminho.rstrip("/")

    itens_brutos = client.list(pasta_raiz, get_info=True, recursive=True)
    raiz_norm = pasta_raiz.strip("/")
    resultado = []
    for item in itens_brutos:
        caminho = caminho_relativo(item)
        if caminho == raiz_norm:
            continue
        resultado.append({
            "path": caminho,
            "etag": item.get("etag"),
            "isdir": bool(item.get("isdir")),
        })
    return resultado


def extrair_pdfs(itens: list[dict], pasta_raiz: str) -> list[dict]:
    """Filtra a árvore completa (listar_arvore) só pros PDFs."""
    return [
        {"path": item["path"], "etag": item["etag"]}
        for item in itens
        # .strip() por causa de arquivo real com espaço sobrando no nome
        # (ex: "Checklist filtro 03.08.26.pdf ") — sem isso, o endswith
        # exato nunca bate e o arquivo some da lista sem nenhum aviso.
        if not item["isdir"] and item["path"].strip().lower().endswith(".pdf")
    ]


def _sem_acento(s: str) -> str:
    return "".join(
        c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)
    ).upper()


def categorizar_equipamento_noel(nome_pasta: str) -> str:
    nome_norm = _sem_acento(nome_pasta)
    for chave, categoria in CATEGORIAS_NOEL:
        if chave in nome_norm:
            return categoria
    return "Outros"


def extrair_loja_noel(caminho_rel: str) -> str | None:
    m = re.search(r"NLJ0*([1-4])", caminho_rel, re.IGNORECASE)
    if m:
        return f"Loja {m.group(1)}"
    m = re.search(r"LOJA\s*0*([1-4])", caminho_rel, re.IGNORECASE)
    if m:
        return f"Loja {m.group(1)}"
    return None


def categoria_equipamento_loja(
    caminho_pdf: str, pasta_raiz: str, cliente: str
) -> tuple[str, str, str | None]:
    rel = caminho_pdf.strip("/")
    raiz_norm = pasta_raiz.strip("/")
    if rel.startswith(raiz_norm):
        rel = rel[len(raiz_norm):].strip("/")
    partes = rel.split("/")

    if cliente == "Mangels":
        categoria = partes[0] if len(partes) >= 1 else "Desconhecida"
        equipamento = partes[-2] if len(partes) >= 2 else "Raiz"
        return categoria, equipamento, None

    # Noel: não confia em profundidade fixa — pega sempre a pasta pai do
    # PDF como equipamento, deriva categoria por palavra-chave e loja pelo
    # prefixo NLJ (ou "LOJA N") em qualquer ponto do caminho.
    equipamento = partes[-2] if len(partes) >= 2 else "Raiz"
    categoria = categorizar_equipamento_noel(equipamento)
    loja = extrair_loja_noel(rel)
    return categoria, equipamento, loja


def extrair_estrutura(itens: list[dict], pasta_raiz: str, cliente: str) -> list[dict]:
    """Lista de equipamentos "no contrato" (pastas), incluindo os que não
    tiveram nenhum atendimento ainda — dá o total real e permite calcular
    pendente = total - realizado. Cada item: {categoria, equipamento, loja}.

    Mangels: só pastas de exatamente 2 níveis (Categoria/Equipamento),
    igual sempre foi. Noel: qualquer pasta cujo nome bata numa categoria
    conhecida vira "equipamento" — não depende de profundidade, e pastas
    organizacionais (ex: a "NOEL LOJA 01" que só embrulha os equipamentos)
    ficam de fora naturalmente por não baterem em nenhuma palavra-chave.
    """
    raiz_norm = pasta_raiz.strip("/")
    resultado: list[dict] = []
    vistos: set[tuple[str, str]] = set()
    for item in itens:
        if not item["isdir"]:
            continue
        caminho = item["path"]
        if not caminho.startswith(raiz_norm + "/"):
            continue
        rel = caminho[len(raiz_norm) + 1:]
        partes = rel.split("/")

        if cliente == "Mangels":
            if len(partes) != 2:
                continue
            categoria, equipamento = partes
            loja = None
        else:
            equipamento = partes[-1]
            categoria = categorizar_equipamento_noel(equipamento)
            if categoria == "Outros":
                continue
            loja = extrair_loja_noel(rel)

        chave = (categoria, equipamento)
        if chave in vistos:
            continue
        vistos.add(chave)
        resultado.append({"categoria": categoria, "equipamento": equipamento, "loja": loja})
    return resultado


def atendimento_para_dict(a: Atendimento) -> dict:
    d = asdict(a)
    d["data"] = a.data.isoformat() if a.data else None
    d["ultima_substituicao_filtro"] = (
        a.ultima_substituicao_filtro.isoformat() if a.ultima_substituicao_filtro else None
    )
    d["proxima_substituicao_filtro"] = (
        a.proxima_substituicao_filtro.isoformat() if a.proxima_substituicao_filtro else None
    )
    return d


def dict_para_atendimento(d: dict) -> Atendimento:
    return Atendimento(
        cliente=d["cliente"],
        categoria=d["categoria"],
        equipamento=d["equipamento"],
        loja=d.get("loja"),
        arquivo=d["arquivo"],
        tipo=d["tipo"],
        data=datetime.fromisoformat(d["data"]) if d["data"] else None,
        duracao_minutos=d["duracao_minutos"],
        ultima_substituicao_filtro=(
            datetime.fromisoformat(d["ultima_substituicao_filtro"])
            if d.get("ultima_substituicao_filtro")
            else None
        ),
        proxima_substituicao_filtro=(
            datetime.fromisoformat(d["proxima_substituicao_filtro"])
            if d.get("proxima_substituicao_filtro")
            else None
        ),
        periodicidade=d.get("periodicidade"),
        vazamento=d.get("vazamento", False),
        gas_kg=d.get("gas_kg"),
    )


def carregar_cache(caminho: str) -> dict:
    if not os.path.exists(caminho):
        return {}
    with open(caminho, "r", encoding="utf-8") as f:
        return json.load(f)


def salvar_cache(cache: dict, caminho: str) -> None:
    with open(caminho, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)


def classificar_tipo(texto: str, categoria: str) -> str:
    # Muitos relatórios do Produttivo têm uma linha só com um ID numérico
    # interno antes do título (ex: "93", "128") — por isso olhamos as
    # primeiras linhas, não só a primeira.
    linhas = [l.strip() for l in texto.splitlines() if l.strip()]
    eh_ordem_servico = False
    for linha in linhas[:3]:
        low = linha.lower()
        if low.startswith("pmoc"):
            return "preventiva"
        if low.startswith("manutenção corretiva") or low.startswith("manutencao corretiva"):
            return "corretiva"
        if low.startswith("check list"):
            # "Check list de filtros bebedouros" — não é preventiva nem
            # corretiva de equipamento, é o registro de troca/vencimento do
            # filtro do bebedouro. Ver extrair_datas_filtro pras datas.
            return "vencimento_filtro"
        if low.startswith("ordem de serviço") or low.startswith("ordem de servico"):
            eh_ordem_servico = True

    if eh_ordem_servico:
        # Chillers ainda não têm template PMOC próprio: até criarem um,
        # "Ordem de Serviço" é o formulário usado pras visitas de rotina.
        if categoria.upper().startswith("CHILLER"):
            return "preventiva"
        return "corretiva"

    return "desconhecida"


def extrair_periodicidade(texto: str) -> str | None:
    """Só confiável quando o título é 'PMOC ... Mensal/Trimestral/Semanal'
    (AR CONDICIONADOS e BEBEDOUROS têm esse template). Chillers não têm
    template de PMOC próprio ainda, então isso sempre volta None pra eles."""
    linhas = [l.strip() for l in texto.splitlines() if l.strip()]
    for linha in linhas[:3]:
        low = linha.lower()
        if not low.startswith("pmoc"):
            continue
        if "trimestral" in low:
            return "trimestral"
        if "semanal" in low:
            return "semanal"
        if "mensal" in low:
            return "mensal"
    return None


def extrair_data(texto: str, nome_arquivo: str) -> datetime | None:
    m = re.search(r"Em:\s*(\d{2})/(\d{2})/(\d{4})", texto)
    if m:
        dia, mes, ano = (int(x) for x in m.groups())
        try:
            return datetime(ano, mes, dia)
        except ValueError:
            pass
    # fallback: data no nome do arquivo (PDF escaneado sem texto, por exemplo)
    m = re.search(r"(\d{2})[.\-](\d{2})[.\-](\d{2,4})", nome_arquivo)
    if m:
        dia, mes, ano = m.groups()
        ano_int = int(ano)
        if ano_int < 100:
            ano_int += 2000
        try:
            return datetime(ano_int, int(mes), int(dia))
        except ValueError:
            return None
    return None


def extrair_duracao_corretiva(texto: str, data_doc: datetime | None) -> float | None:
    m_inicio = re.search(
        r"Início do Trabalho.*?Horário de início\s*\n?\s*(\d{2}:\d{2}:\d{2})",
        texto,
        re.DOTALL,
    )
    m_fim = re.search(
        r"Fim do Trabalho.*?Horário\s*\n?\s*(\d{2}/\d{2}/\d{4})\s+(\d{2}:\d{2}:\d{2})",
        texto,
        re.DOTALL,
    )
    if not m_inicio or not m_fim:
        return None
    fim_dt = datetime.strptime(f"{m_fim.group(1)} {m_fim.group(2)}", "%d/%m/%Y %H:%M:%S")
    inicio_time = datetime.strptime(m_inicio.group(1), "%H:%M:%S").time()
    inicio_dt = datetime.combine(fim_dt.date(), inicio_time)
    if inicio_dt > fim_dt:
        inicio_dt -= timedelta(days=1)
    return (fim_dt - inicio_dt).total_seconds() / 60


def extrair_datas_filtro(texto: str) -> tuple[datetime | None, datetime | None]:
    """Só o relatório 'Check list de filtros bebedouros' tem esses campos.
    Retorna (última substituição, próxima substituição)."""
    m = re.search(
        r"Data da (?:ultima|última) substi\w*[cç][aã]o\s+Data da pr[oó]xima substi\w*[cç][aã]o"
        r"\s*\n?\s*(\d{2})/(\d{2})/(\d{4})[^\n]*?(\d{2})/(\d{2})/(\d{4})",
        texto,
    )
    if not m:
        return None, None
    d1, m1, a1, d2, m2, a2 = m.groups()
    try:
        ultima = datetime(int(a1), int(m1), int(d1))
    except ValueError:
        ultima = None
    try:
        proxima = datetime(int(a2), int(m2), int(d2))
    except ValueError:
        proxima = None
    return ultima, proxima


def detectar_vazamento(texto: str, tipo: str) -> bool:
    # Só em corretivas: os checklists de PMOC/preventiva têm "verificar
    # vazamento" como item de rotina do formulário, então a palavra aparece
    # ali mesmo sem vazamento nenhum — contar isso inflava a recorrência
    # com falso positivo (confirmado em 06/08/2026, rodando contra dados
    # reais do Noel: 4 de 5 relatórios de um equipamento vinham marcados
    # como vazamento, sendo 4 preventivas de rotina e só 1 vazamento real).
    if tipo != "corretiva":
        return False
    return "vazamento" in texto.lower()


def extrair_gas_kg(texto: str) -> float | None:
    """Tenta achar uma quantidade de gás em kg no texto (Diagnóstico/
    Solução/Observações). Hoje os relatórios do Noel não têm esse campo
    estruturado (confirmado com o Rafael em 06/08/2026) — retorna None
    até o time começar a registrar isso, o que não é um erro."""
    m = re.search(r"(\d+[.,]?\d*)\s*kg", texto, re.IGNORECASE)
    if not m:
        return None
    try:
        return float(m.group(1).replace(",", "."))
    except ValueError:
        return None


def processar_pdf(client: Client, caminho_remoto: str, pasta_raiz: str, cliente: str) -> Atendimento:
    categoria, equipamento, loja = categoria_equipamento_loja(caminho_remoto, pasta_raiz, cliente)
    nome_arquivo = caminho_remoto.rstrip("/").split("/")[-1]
    caminho_local = os.path.join(PASTA_TEMP, nome_arquivo)

    client.download_sync(remote_path=caminho_remoto, local_path=caminho_local)

    texto = ""
    try:
        with pdfplumber.open(caminho_local) as pdf:
            texto = "\n".join((p.extract_text() or "") for p in pdf.pages)
    finally:
        if os.path.exists(caminho_local):
            os.remove(caminho_local)

    tipo = classificar_tipo(texto, categoria)
    data = extrair_data(texto, nome_arquivo)
    duracao = extrair_duracao_corretiva(texto, data) if tipo == "corretiva" else None
    ultima_filtro, proxima_filtro = extrair_datas_filtro(texto)
    periodicidade = extrair_periodicidade(texto) if tipo == "preventiva" else None
    vazamento = detectar_vazamento(texto, tipo)
    gas_kg = extrair_gas_kg(texto) if vazamento else None

    if tipo == "vencimento_filtro" and data and proxima_filtro and proxima_filtro.date() == data.date():
        # Bug do formulário Produttivo: quando o campo "próxima substituição"
        # não é preenchido pelo técnico, ele volta com a mesma data/hora do
        # preenchimento do checklist (não é uma substituição vencida hoje —
        # é a visita de hoje que gera este checklist). Nesse caso, a data
        # real de referência é a própria visita: última substituição = hoje,
        # próxima = hoje + periodicidade.
        ultima_filtro = datetime(data.year, data.month, data.day)
        proxima_filtro = ultima_filtro + timedelta(days=PERIODICIDADE_FILTRO_DIAS)

    return Atendimento(
        cliente=cliente,
        categoria=categoria,
        equipamento=equipamento,
        loja=loja,
        arquivo=nome_arquivo,
        tipo=tipo,
        data=data,
        duracao_minutos=duracao,
        ultima_substituicao_filtro=ultima_filtro,
        proxima_substituicao_filtro=proxima_filtro,
        periodicidade=periodicidade,
        vazamento=vazamento,
        gas_kg=gas_kg,
    )


def gerar_relatorio(atendimentos: list[Atendimento], caminho_saida: str, titulo: str) -> None:
    doc = Document()
    doc.add_heading(titulo, level=0)
    doc.add_paragraph(f"Gerado em {datetime.now():%d/%m/%Y %H:%M} — Ano de referência: 2026")

    validos = [a for a in atendimentos if a.tipo != "desconhecida"]
    nao_classificados = [a for a in atendimentos if a.tipo == "desconhecida"]
    preventivas = [a for a in validos if a.tipo == "preventiva"]
    corretivas = [a for a in validos if a.tipo == "corretiva"]
    horas_corretivas_total = sum(
        (a.duracao_minutos or 0) for a in corretivas
    ) / 60

    doc.add_heading("Resumo Geral", level=1)
    resumo = doc.add_paragraph()
    resumo.add_run(f"Total de atendimentos encontrados: {len(atendimentos)}\n")
    resumo.add_run(f"Preventivas: {len(preventivas)}\n")
    resumo.add_run(f"Corretivas: {len(corretivas)}\n")
    resumo.add_run(f"Total de horas corretivas: {horas_corretivas_total:.1f}h\n")
    if nao_classificados:
        resumo.add_run(
            f"⚠ {len(nao_classificados)} arquivo(s) não classificado(s) — ver apêndice ao final.\n"
        )

    categorias = sorted({a.categoria for a in validos})
    doc.add_heading("Detalhamento por Categoria e Equipamento", level=1)
    for categoria in categorias:
        doc.add_heading(categoria, level=2)
        do_cat = [a for a in validos if a.categoria == categoria]
        equipamentos = sorted({a.equipamento for a in do_cat})

        linhas = []
        for equipamento in equipamentos:
            do_equip = [a for a in do_cat if a.equipamento == equipamento]
            n_prev = sum(1 for a in do_equip if a.tipo == "preventiva")
            n_corr = sum(1 for a in do_equip if a.tipo == "corretiva")
            horas = sum((a.duracao_minutos or 0) for a in do_equip if a.tipo == "corretiva") / 60
            linhas.append((equipamento, n_prev, n_corr, horas))

        # equipamentos mais problemáticos (mais horas corretivas) primeiro
        linhas.sort(key=lambda r: r[3], reverse=True)

        tabela = doc.add_table(rows=1, cols=4)
        tabela.style = "Light Grid Accent 1"
        hdr = tabela.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            "Equipamento",
            "Preventivas",
            "Corretivas",
            "Horas Corretivas",
        )
        for equipamento, n_prev, n_corr, horas in linhas:
            row = tabela.add_row().cells
            row[0].text = equipamento
            row[1].text = str(n_prev)
            row[2].text = str(n_corr)
            row[3].text = f"{horas:.1f}h"

    if nao_classificados:
        doc.add_heading("Apêndice — Arquivos Não Classificados", level=1)
        doc.add_paragraph(
            "Estes arquivos não começam com 'PMOC' nem 'Ordem de Serviço' no "
            "texto extraído (pode ser PDF escaneado sem texto, ou um formato "
            "de relatório diferente). Revisar manualmente:"
        )
        for a in nao_classificados:
            doc.add_paragraph(f"{a.categoria} / {a.equipamento} / {a.arquivo}", style="List Bullet")

    doc.save(caminho_saida)


def enviar_para_supabase(pares: list[tuple[str, Atendimento]], cliente: str) -> None:
    """Substitui o conteúdo de kpi_manutencao_atendimentos DESTE CLIENTE
    pelo resultado desta execução — o DELETE é escopado por `cliente`, não
    apaga as linhas dos outros clientes. Dataset é pequeno (algumas
    centenas de linhas) — mais simples que fazer upsert incremental linha
    a linha."""
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": "application/json",
    }
    tabela_url = f"{SUPABASE_URL}/rest/v1/kpi_manutencao_atendimentos"

    resp = requests.delete(f"{tabela_url}?cliente=eq.{cliente}", headers=headers)
    resp.raise_for_status()

    registros = [
        {
            "caminho": caminho,
            "cliente": a.cliente,
            "categoria": a.categoria,
            "equipamento": a.equipamento,
            "loja": a.loja,
            "arquivo": a.arquivo,
            "tipo": a.tipo,
            "data": a.data.date().isoformat() if a.data else None,
            "duracao_minutos": a.duracao_minutos,
            "vazamento": a.vazamento,
            "gas_kg": a.gas_kg,
        }
        for caminho, a in pares
    ]

    # insere em lotes pra não estourar o tamanho do payload
    tamanho_lote = 200
    for i in range(0, len(registros), tamanho_lote):
        lote = registros[i : i + tamanho_lote]
        resp = requests.post(tabela_url, headers=headers, json=lote)
        resp.raise_for_status()


def montar_resumo_equipamentos(
    estrutura: list[dict],
    pares: list[tuple[str, Atendimento]],
    data_referencia: datetime | None = None,
) -> list[dict]:
    """Um registro por equipamento (pasta), mesmo os que não tiveram
    nenhum atendimento ainda — é o que dá o "total no contrato" e permite
    calcular pendente = total - realizado.

    `data_referencia` define o "mês atual" usado pra calcular
    realizado/pendente — por padrão é hoje, mas pode ser simulado (ex:
    --mes-referencia 2026-08-01) só pra pré-visualizar o painel."""
    agora = data_referencia or datetime.now()
    por_equip: dict[tuple[str, str], dict] = {}

    def registro_vazio(categoria: str, equipamento: str, loja: str | None) -> dict:
        return {
            "categoria": categoria,
            "equipamento": equipamento,
            "loja": loja,
            "ultima_preventiva": None,
            "periodicidade": None,
            "ultima_substituicao_filtro": None,
            "proxima_substituicao_filtro": None,
            # data do checklist de filtro mais recente já visto — usada só
            # internamente pra saber qual "próxima substituição" é a mais
            # atual (a de um checklist antigo pode já estar desatualizada).
            "_data_checklist_filtro": None,
            # "teve preventiva neste mês calendário?" — critério de em_dia
            # pra todas as categorias (ver decisão de 03/08/2026 acima).
            "_realizado_mes_atual": False,
        }

    for e in estrutura:
        por_equip[(e["categoria"], e["equipamento"])] = registro_vazio(
            e["categoria"], e["equipamento"], e["loja"]
        )

    for _, a in pares:
        chave = (a.categoria, a.equipamento)
        if chave not in por_equip:
            # atendimento cuja pasta não apareceu na listagem de estrutura
            # (não deveria acontecer, mas não descartamos o dado por isso)
            por_equip[chave] = registro_vazio(a.categoria, a.equipamento, a.loja)
        info = por_equip[chave]
        if a.tipo == "preventiva" and a.data:
            if info["ultima_preventiva"] is None or a.data > info["ultima_preventiva"]:
                info["ultima_preventiva"] = a.data
                info["periodicidade"] = a.periodicidade
            if a.data.year == agora.year and a.data.month == agora.month:
                info["_realizado_mes_atual"] = True
        if a.proxima_substituicao_filtro and a.data:
            if info["_data_checklist_filtro"] is None or a.data > info["_data_checklist_filtro"]:
                info["_data_checklist_filtro"] = a.data
                info["ultima_substituicao_filtro"] = a.ultima_substituicao_filtro
                info["proxima_substituicao_filtro"] = a.proxima_substituicao_filtro

    # Decidido com o Rafael em 03/08/2026: voltar pra regra simples de mês
    # calendário pra TODAS as categorias (não só chillers) — "em dia" só se
    # teve preventiva dentro do mês vigente, sem carência pela periodicidade
    # própria (mensal/trimestral/semanal). Reseta pra 0 feito no dia 1 de
    # cada mês, mesmo que o prazo real do PMOC ainda não tenha vencido.
    for info in por_equip.values():
        info["em_dia"] = info["_realizado_mes_atual"]
        info["dias_atraso"] = None
        del info["_realizado_mes_atual"]

    return list(por_equip.values())


def enviar_equipamentos_para_supabase(registros: list[dict], cliente: str) -> None:
    headers = {
        "apikey": SUPABASE_SERVICE_ROLE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_ROLE_KEY}",
        "Content-Type": "application/json",
    }
    tabela_url = f"{SUPABASE_URL}/rest/v1/kpi_manutencao_equipamentos"

    resp = requests.delete(f"{tabela_url}?cliente=eq.{cliente}", headers=headers)
    resp.raise_for_status()

    payload = [
        {
            "cliente": cliente,
            "categoria": r["categoria"],
            "equipamento": r["equipamento"],
            "loja": r["loja"],
            "em_dia": r["em_dia"],
            "dias_atraso": r["dias_atraso"],
            "periodicidade": r["periodicidade"],
            "ultima_preventiva": r["ultima_preventiva"].date().isoformat()
            if r["ultima_preventiva"]
            else None,
            "ultima_substituicao_filtro": r["ultima_substituicao_filtro"].date().isoformat()
            if r["ultima_substituicao_filtro"]
            else None,
            "proxima_substituicao_filtro": r["proxima_substituicao_filtro"].date().isoformat()
            if r["proxima_substituicao_filtro"]
            else None,
        }
        for r in registros
    ]

    tamanho_lote = 200
    for i in range(0, len(payload), tamanho_lote):
        lote = payload[i : i + tamanho_lote]
        resp = requests.post(tabela_url, headers=headers, json=lote)
        resp.raise_for_status()


def processar_cliente(
    client: Client,
    nome_cliente: str,
    pasta_raiz: str,
    cache_path: str,
    data_referencia: datetime | None,
) -> None:
    max_novos = int(os.environ.get("KPI_MAX_NOVOS_POR_EXECUCAO", "0") or "0")

    cache_anterior = carregar_cache(cache_path)

    print(f"Listando '{pasta_raiz}' (uma chamada recursiva só)...")
    arvore = listar_arvore(client, pasta_raiz)
    pdfs = extrair_pdfs(arvore, pasta_raiz)
    estrutura = extrair_estrutura(arvore, pasta_raiz, nome_cliente)
    print(f"{len(pdfs)} PDF(s) encontrado(s), {len(estrutura)} equipamento(s) na estrutura.")

    atendimentos: list[Atendimento] = []
    pares: list[tuple[str, Atendimento]] = []
    cache_novo: dict = dict(cache_anterior)
    reaproveitados = 0
    novos_processados = 0
    execucao_completa = True
    for i, item in enumerate(pdfs, 1):
        caminho, etag = item["path"], item["etag"]
        nome = caminho.rstrip("/").split("/")[-1]
        entrada_cache = cache_anterior.get(caminho)

        if (
            entrada_cache
            and entrada_cache.get("etag") == etag
            and entrada_cache.get("versao") == CACHE_VERSAO
        ):
            a = dict_para_atendimento(entrada_cache["dados"])
            atendimentos.append(a)
            pares.append((caminho, a))
            reaproveitados += 1
            continue

        if max_novos and novos_processados >= max_novos:
            execucao_completa = False
            print(
                f"\nLimite de {max_novos} novo(s)/alterado(s) atingido "
                f"({i - 1}/{len(pdfs)} arquivos vistos) — parando por aqui, "
                "a próxima execução continua de onde parou."
            )
            break

        print(f"  [{i}/{len(pdfs)}] {nome} (novo/alterado/regra atualizada)")
        try:
            a = processar_pdf(client, caminho, pasta_raiz, nome_cliente)
            atendimentos.append(a)
            pares.append((caminho, a))
            cache_novo[caminho] = {
                "etag": etag,
                "versao": CACHE_VERSAO,
                "dados": atendimento_para_dict(a),
            }
            novos_processados += 1
        except Exception as e:
            print(f"    [erro] não processado: {e}")

    salvar_cache(cache_novo, cache_path)

    print(f"\n{novos_processados} processado(s) agora, {reaproveitados} reaproveitado(s) do cache.")

    if not execucao_completa:
        print(
            "Execução parcial — não publica no Supabase ainda (lista não foi "
            "percorrida por inteiro). Rode de novo pra continuar o lote."
        )
        return

    saida_docx = f"kpi_manutencao_{nome_cliente.lower()}_{datetime.now():%Y%m%d}.docx"
    print(f"Gerando relatório: {saida_docx}")
    gerar_relatorio(atendimentos, saida_docx, f"KPI de Manutenção — {nome_cliente}")

    print("Publicando no Supabase...")
    enviar_para_supabase(pares, nome_cliente)

    resumo_equipamentos = montar_resumo_equipamentos(estrutura, pares, data_referencia)
    enviar_equipamentos_para_supabase(resumo_equipamentos, nome_cliente)


def main():
    import argparse

    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--mes-referencia",
        help="Simula outra data pra calcular realizado/pendente do mês "
        "(formato AAAA-MM-DD). Sem isso, usa a data de hoje de verdade — "
        "só usar pra pré-visualizar o painel, depois rodar de novo sem a "
        "flag pra restaurar os dados reais.",
    )
    parser.add_argument(
        "--cliente",
        help="Roda só um cliente (ex: Noel) em vez de todos os configurados "
        "em CLIENTES — útil pra testar isolado sem re-processar tudo.",
    )
    args = parser.parse_args()
    data_referencia = (
        datetime.strptime(args.mes_referencia, "%Y-%m-%d") if args.mes_referencia else None
    )
    if data_referencia:
        print(f"⚠ SIMULANDO data de referência: {data_referencia:%d/%m/%Y} (não é a data real)")

    clientes_a_rodar = CLIENTES
    if args.cliente:
        clientes_a_rodar = [c for c in CLIENTES if c["nome"].lower() == args.cliente.lower()]
        if not clientes_a_rodar:
            nomes = ", ".join(c["nome"] for c in CLIENTES)
            print(f"Cliente '{args.cliente}' não configurado. Opções: {nomes}")
            return

    if os.path.exists(PASTA_TEMP):
        shutil.rmtree(PASTA_TEMP)
    os.makedirs(PASTA_TEMP, exist_ok=True)

    client = conectar_nextcloud()

    for config in clientes_a_rodar:
        print(f"\n=== {config['nome']} ===")
        try:
            processar_cliente(
                client, config["nome"], config["pasta_raiz"], config["cache_path"], data_referencia
            )
        except Exception as e:
            # Um cliente com problema não pode travar os outros — cada um
            # publica (ou não) de forma independente.
            print(f"[erro] Falha ao processar {config['nome']}: {e}")

    shutil.rmtree(PASTA_TEMP, ignore_errors=True)
    print("\nConcluído.")


if __name__ == "__main__":
    main()
